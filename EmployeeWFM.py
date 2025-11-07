import os
import sys
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import urllib.request
import urllib.parse
import json
import threading
import time
import io
import datetime

def is_pyinstaller_bundle():
    """Check if running as PyInstaller bundle"""
    return getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS')

class PerplexityAI:
    # API key as class constant - update this with your actual API key
    API_KEY = "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"
    
    def __init__(self, api_key=None):
        self.api_key = api_key or self.API_KEY
        self.base_url = "https://api.perplexity.ai/chat/completions"
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
    
    def query(self, prompt, model="sonar", disable_search=False):
        """
        Send a query to Perplexity AI and return the response
        
        Args:
            prompt (str): The question/prompt to send
            model (str): The model to use (default: sonar - cheapest and fastest)
            disable_search (bool): If True, disables web search and uses only training data
        
        Returns:
            str: The AI response or error message
        """
        try:
            # Prepare the request data
            data = {
                "model": model,
                "messages": [
                    {
                        "role": "system",
                        "content": "You are a helpful AI assistant. Provide clear, accurate, and concise responses."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                "max_tokens": 2000,  # Increased for longer document analysis
                "temperature": 0.2,
                "top_p": 0.9,
                "stream": False
            }
            
            # Add disable_search parameter if requested
            if disable_search:
                data["disable_search"] = True
            
            # Convert data to JSON
            json_data = json.dumps(data).encode('utf-8')
            
            # Create the request
            request = urllib.request.Request(
                self.base_url,
                data=json_data,
                headers=self.headers,
                method='POST'
            )
            
            # Send the request
            with urllib.request.urlopen(request) as response:
                response_data = json.loads(response.read().decode('utf-8'))
                
                # Extract the AI response
                if 'choices' in response_data and len(response_data['choices']) > 0:
                    return response_data['choices'][0]['message']['content']
                else:
                    return "❌ No response received from AI"
                    
        except urllib.error.HTTPError as e:
            error_body = e.read().decode('utf-8')
            try:
                error_json = json.loads(error_body)
                error_msg = error_json.get('error', {}).get('message', 'Unknown error')
                return f"❌ HTTP Error {e.code}: {error_msg}"
            except:
                return f"❌ HTTP Error {e.code}: {error_body}"
                
        except urllib.error.URLError as e:
            return f"❌ Network Error: {e.reason}"
            
        except json.JSONDecodeError as e:
            return f"❌ JSON Error: {e}"
            
        except Exception as e:
            return f"❌ Unexpected Error: {e}"

class TimerDialog:
    def __init__(self, parent, title="Processing", message="Analyzing files with AI..."):
        self.parent = parent
        self.title = title
        self.message = message
        self.elapsed_time = 0
        self.timer_job = None
        self.dialog = None
        self.time_label = None
        
    def show(self):
        """Show the timer dialog"""
        # Create the dialog window
        self.dialog = tk.Toplevel(self.parent)
        self.dialog.title(self.title)
        self.dialog.configure(bg='#f0f0f0')
        self.dialog.geometry("400x150")
        self.dialog.resizable(False, False)
        
        # Center the window
        self.dialog.transient(self.parent)
        self.dialog.grab_set()
        
        # Center on screen
        x = (self.dialog.winfo_screenwidth() // 2) - 200
        y = (self.dialog.winfo_screenheight() // 2) - 75
        self.dialog.geometry(f"400x150+{x}+{y}")
        
        # Message label
        message_label = tk.Label(
            self.dialog,
            text=self.message,
            font=("Arial", 12),
            bg='#f0f0f0',
            fg='#333333',
            wraplength=350
        )
        message_label.pack(pady=(20, 10))
        
        # Timer label
        self.time_label = tk.Label(
            self.dialog,
            text="Elapsed time: 0 seconds",
            font=("Arial", 14, "bold"),
            bg='#f0f0f0',
            fg='#2196F3'
        )
        self.time_label.pack(pady=10)
        
        # Progress indicator (animated dots)
        self.dots_label = tk.Label(
            self.dialog,
            text="●",
            font=("Arial", 16),
            bg='#f0f0f0',
            fg='#4CAF50'
        )
        self.dots_label.pack(pady=10)
        
        # Start the timer
        self.start_timer()
        
        # Start dots animation
        self.animate_dots()
        
        return self.dialog
    
    def start_timer(self):
        """Start the timer that counts up every second"""
        if self.dialog and self.dialog.winfo_exists():
            self.elapsed_time += 1
            if self.time_label and self.time_label.winfo_exists():
                self.time_label.config(text=f"Elapsed time: {self.elapsed_time} seconds")
            
            # Update the dialog
            self.dialog.update()
            
            # Schedule next update
            self.timer_job = self.dialog.after(1000, self.start_timer)
    
    def animate_dots(self):
        """Animate the dots for visual feedback"""
        if self.dots_label and self.dialog and self.dialog.winfo_exists():
            try:
                current_dots = self.dots_label.cget("text")
                if current_dots == "●":
                    new_dots = "●●"
                elif current_dots == "●●":
                    new_dots = "●●●"
                else:
                    new_dots = "●"
                
                self.dots_label.config(text=new_dots)
                self.dialog.update()
                self.dialog.after(500, self.animate_dots)
            except tk.TclError:
                pass  # Dialog was destroyed
    
    def close(self):
        """Close the timer dialog"""
        if self.timer_job:
            self.dialog.after_cancel(self.timer_job)
        if self.dialog and self.dialog.winfo_exists():
            self.dialog.destroy()

class RoundedButton:
    def __init__(self, parent, text, command=None, bg='#4CAF50', fg='white', 
                 font=('Arial', 10), padx=20, pady=10, corner_radius=10, 
                 hover_bg=None, cursor='hand2'):
        self.parent = parent
        self.text = text
        self.command = command
        self.bg = bg
        self.fg = fg
        self.font = font
        self.padx = padx
        self.pady = pady
        self.corner_radius = corner_radius
        self.hover_bg = hover_bg if hover_bg else self._darken_color(bg)
        self.cursor = cursor
        self.is_hovered = False
        
        # Calculate button dimensions
        temp_label = tk.Label(parent, text=text, font=font)
        temp_label.update_idletasks()
        text_width = temp_label.winfo_reqwidth()
        text_height = temp_label.winfo_reqheight()
        temp_label.destroy()
        
        self.width = text_width + (padx * 2)
        self.height = text_height + (pady * 2)
        
        # Create canvas
        self.canvas = tk.Canvas(
            parent, 
            width=self.width, 
            height=self.height,
            highlightthickness=0,
            cursor=cursor
        )
        
        self.draw_button()
        self.bind_events()
    
    def _darken_color(self, color):
        """Darken a hex color for hover effect"""
        if color.startswith('#'):
            color = color[1:]
        try:
            r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
            r = max(0, r - 30)
            g = max(0, g - 30)
            b = max(0, b - 30)
            return f"#{r:02x}{g:02x}{b:02x}"
        except:
            return '#333333'
    
    def draw_button(self):
        """Draw the rounded rectangle button"""
        self.canvas.delete("all")
        
        # Choose color based on hover state
        current_bg = self.hover_bg if self.is_hovered else self.bg
        
        # Draw rounded rectangle
        x1, y1 = 0, 0
        x2, y2 = self.width, self.height
        r = self.corner_radius
        
        # Create rounded rectangle using multiple shapes
        self.canvas.create_arc(x1, y1, x1 + 2*r, y1 + 2*r, start=90, extent=90, 
                              fill=current_bg, outline=current_bg)
        self.canvas.create_arc(x2 - 2*r, y1, x2, y1 + 2*r, start=0, extent=90, 
                              fill=current_bg, outline=current_bg)
        self.canvas.create_arc(x1, y2 - 2*r, x1 + 2*r, y2, start=180, extent=90, 
                              fill=current_bg, outline=current_bg)
        self.canvas.create_arc(x2 - 2*r, y2 - 2*r, x2, y2, start=270, extent=90, 
                              fill=current_bg, outline=current_bg)
        
        # Fill the middle areas
        self.canvas.create_rectangle(x1 + r, y1, x2 - r, y2, 
                                   fill=current_bg, outline=current_bg)
        self.canvas.create_rectangle(x1, y1 + r, x2, y2 - r, 
                                   fill=current_bg, outline=current_bg)
        
        # Add text
        self.canvas.create_text(self.width//2, self.height//2, text=self.text, 
                               fill=self.fg, font=self.font, anchor='center')
    
    def bind_events(self):
        """Bind mouse events"""
        self.canvas.bind("<Button-1>", self.on_click)
        self.canvas.bind("<Enter>", self.on_enter)
        self.canvas.bind("<Leave>", self.on_leave)
    
    def on_click(self, event):
        """Handle button click"""
        if self.command:
            self.command()
    
    def on_enter(self, event):
        """Handle mouse enter"""
        self.is_hovered = True
        self.draw_button()
    
    def on_leave(self, event):
        """Handle mouse leave"""
        self.is_hovered = False
        self.draw_button()
    
    def pack(self, **kwargs):
        """Pack the canvas"""
        return self.canvas.pack(**kwargs)
    
    def grid(self, **kwargs):
        """Grid the canvas"""
        return self.canvas.grid(**kwargs)
    
    def place(self, **kwargs):
        """Place the canvas"""
        return self.canvas.place(**kwargs)
    
    def config(self, **kwargs):
        """Configure button properties"""
        if 'text' in kwargs:
            self.text = kwargs['text']
        if 'bg' in kwargs:
            self.bg = kwargs['bg']
        if 'fg' in kwargs:
            self.fg = kwargs['fg']
        self.draw_button()
    
    def destroy(self):
        """Destroy the button"""
        self.canvas.destroy()

def read_file_as_variable(file_path):
    """
    Reads a file (CSV, PDF, or Word) and returns its content as a variable.
    Does not decode the file - returns raw content.
    
    Args:
        file_path (str): Path to the file to read
        
    Returns:
        bytes: Raw file content as bytes
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file type is not supported
    """
    
    # Convert to Path object for easier handling
    file_path = Path(file_path)
    
    # Check if file exists
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Get file extension
    file_extension = file_path.suffix.lower()
    
    # Supported file types
    supported_extensions = ['.csv', '.pdf', '.doc', '.docx']
    
    if file_extension not in supported_extensions:
        raise ValueError(f"Unsupported file type: {file_extension}. Supported types: {supported_extensions}")
    
    # Read file as binary (raw bytes)
    with open(file_path, 'rb') as file:
        file_content = file.read()
    
    return file_content

def process_multiple_files(file_paths):
    """
    Process multiple files and return a dictionary with file names as keys
    and file content as values.
    
    Args:
        file_paths (list): List of file paths to process
        
    Returns:
        dict: Dictionary with filename as key and file content as value
    """
    
    file_variables = {}
    
    for file_path in file_paths:
        try:
            file_name = Path(file_path).name
            file_content = read_file_as_variable(file_path)
            file_variables[file_name] = file_content
            print(f"Successfully loaded: {file_name} ({len(file_content)} bytes)")
        except Exception as e:
            print(f"Error loading {file_path}: {e}")
    
    return file_variables

def get_file_info(file_variable):
    """
    Get basic information about the file variable.
    
    Args:
        file_variable (bytes): The file content as bytes
        
    Returns:
        dict: File information
    """
    return {
        'size_bytes': len(file_variable),
        'size_kb': len(file_variable) / 1024,
        'size_mb': len(file_variable) / (1024 * 1024),
        'type': type(file_variable).__name__
    }

def save_file_variable(file_variable, output_path):
    """
    Save a file variable back to disk.
    
    Args:
        file_variable (bytes): The file content as bytes
        output_path (str): Path where to save the file
    """
    with open(output_path, 'wb') as file:
        file.write(file_variable)
    print(f"File saved to: {output_path}")

class FileUploaderGUI:
    #GUI SETUP & LAYOUT FUNCTIONS
    def __init__(self, root):
        self.root = root
        self.root.title("Employee WFM - File Uploader")
        self.root.geometry("600x400")
        self.root.configure(bg='#f0f0f0')
        
        # Store uploaded files
        self.uploaded_files = {}
        
        # Store output files (generated from queries)
        self.output_files = {}
        
        self.create_widgets()
    
    def create_widgets(self):
        # Title
        title_label = tk.Label(
            self.root, 
            text="Employee WFM Audit", 
            font=("Arial", 16, "bold"),
            bg='#f0f0f0',
            fg='#333333'
        )
        title_label.pack(pady=20)
        
        # Policy subtitle (clickable button)
        policy_btn = RoundedButton(
            self.root,
            text="Open Work From Home Policy",
            command=self.open_policy_pdf,
            bg='#e3f2fd',
            fg='#1976d2',
            hover_bg='#bbdefb',
            font=("Arial", 10),
            padx=20,
            pady=6,
            corner_radius=10
        )
        policy_btn.pack(pady=(0, 5))
        
        # Instructions
        instructions = tk.Label(
            self.root,
            text="Click a button below to upload CSV, PDF, or Word files",
            font=("Arial", 10),
            bg='#f0f0f0',
            fg='#666666'
        )
        instructions.pack(pady=3)
        
        # Upload buttons frame
        upload_frame = tk.Frame(self.root, bg='#f0f0f0')
        upload_frame.pack(pady=10)
        
        # Upload button
        upload_btn = RoundedButton(
            upload_frame,
            text="📁 Upload File",
            command=self.upload_file,
            bg='#2196F3',
            fg='white',
            font=("Arial", 10),
            padx=20,
            pady=6,
            corner_radius=10
        )
        upload_btn.pack(side='left', padx=10)
        
        # Upload multiple files button
        upload_multiple_btn = RoundedButton(
            upload_frame,
            text="📁 Upload Folder",
            command=self.upload_folder,
            bg='#2196F3',
            fg='white',
            font=("Arial", 10),
            padx=20,
            pady=6,
            corner_radius=10
        )
        upload_multiple_btn.pack(side='left', padx=10)
        
        # Summary button (below upload buttons)
        summary_btn = RoundedButton(
            self.root,
            text="Create Summary",
            command=self.make_summary,
            bg='#FF9800',
            fg='white',
            font=("Arial", 10),
            padx=20,
            pady=6,
            corner_radius=10
        )
        summary_btn.pack(pady=5)
        
        # Main content frame for two columns
        main_frame = tk.Frame(self.root, bg='#f0f0f0')
        main_frame.pack(pady=10, padx=20, fill='both', expand=True)
        
        # Configure main_frame for proportional columns (2/3 and 1/3)
        main_frame.grid_rowconfigure(0, weight=1)
        main_frame.grid_columnconfigure(0, weight=2)  # 2/3 for uploaded files
        main_frame.grid_columnconfigure(1, weight=1)  # 1/3 for output
        
        # Left column - Uploaded Files (2/3 of width)
        left_frame = tk.Frame(main_frame, bg='#f0f0f0')
        left_frame.grid(row=0, column=0, sticky='nsew', padx=(0, 5))
        
        # File list label
        list_label = tk.Label(
            left_frame,
            text="Uploaded Files:",
            font=("Arial", 12, "bold"),
            bg='#f0f0f0',
            fg='#333333'
        )
        list_label.pack(anchor='w')
        
        # Scrollable frame for file list with query buttons
        self.canvas = tk.Canvas(left_frame, bg='white')
        self.scrollbar = tk.Scrollbar(left_frame, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = tk.Frame(self.canvas, bg='white')
        
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        self.canvas.pack(side="left", fill="both", expand=True, pady=5)
        self.scrollbar.pack(side="right", fill="y", pady=5)
        
        # Right column - Output Files (1/3 of width)
        right_frame = tk.Frame(main_frame, bg='#f0f0f0')
        right_frame.grid(row=0, column=1, sticky='nsew', padx=(5, 0))
        
        # Output list label
        output_label = tk.Label(
            right_frame,
            text="Output:",
            font=("Arial", 12, "bold"),
            bg='#f0f0f0',
            fg='#333333'
        )
        output_label.pack(anchor='w')
        
        # Scrollable frame for output files
        self.output_canvas = tk.Canvas(right_frame, bg='white')
        self.output_scrollbar = tk.Scrollbar(right_frame, orient="vertical", command=self.output_canvas.yview)
        self.output_scrollable_frame = tk.Frame(self.output_canvas, bg='white')
        
        self.output_scrollable_frame.bind(
            "<Configure>",
            lambda e: self.output_canvas.configure(scrollregion=self.output_canvas.bbox("all"))
        )
        
        self.output_canvas.create_window((0, 0), window=self.output_scrollable_frame, anchor="nw")
        self.output_canvas.configure(yscrollcommand=self.output_scrollbar.set)
        
        self.output_canvas.pack(side="left", fill="both", expand=True, pady=5)
        self.output_scrollbar.pack(side="right", fill="y", pady=5)
        
        # Buttons frame
        buttons_frame = tk.Frame(self.root, bg='#f0f0f0')
        buttons_frame.pack(pady=10)
        
        # Clear all button
        clear_btn = RoundedButton(
            buttons_frame,
            text="Clear All Uploaded",
            command=self.clear_all_files,
            bg='#f44336',
            fg='white',
            font=("Arial", 10),
            padx=20,
            pady=5,
            corner_radius=10
        )
        clear_btn.pack(side='left', padx=5)
        
        # Clear all outputs button
        clear_outputs_btn = RoundedButton(
            buttons_frame,
            text="Clear All Outputs",
            command=self.clear_all_outputs,
            bg='#f44336',
            fg='white',
            font=("Arial", 10),
            padx=20,
            pady=5,
            corner_radius=10
        )
        clear_outputs_btn.pack(side='left', padx=5)
    
    def update_file_list(self):
        """Update the file list display with query buttons"""
        # Clear existing widgets in scrollable frame
        for widget in self.scrollable_frame.winfo_children():
            widget.destroy()
        
        for i, (filename, file_data) in enumerate(self.uploaded_files.items()):
            content = file_data['content']
            size_kb = len(content) / 1024
            if size_kb < 1024:
                size_str = f"{size_kb:.1f} KB"
            else:
                size_str = f"{size_kb/1024:.1f} MB"
            
            # Create frame for each file entry - transparent background
            file_frame = tk.Frame(self.scrollable_frame, bg='white', relief='flat', bd=0, highlightthickness=0)
            file_frame.pack(fill='x', padx=3, pady=1)
            
            # File info label (clickable to open file)
            file_info_label = RoundedButton(
                file_frame,
                text=f"📄 {filename} ({size_str})",
                command=lambda f=filename: self.open_file(f),
                bg='#f8f9fa',
                fg='#333333',
                hover_bg='#e9ecef',
                font=("Arial", 10),
                padx=12,
                pady=8,
                corner_radius=6
            )
            file_info_label.pack(side='left', fill='x', expand=True, padx=(0, 1), pady=0)
            
            # Red X button to remove file
            remove_btn = RoundedButton(
                file_frame,
                text="✕",
                command=lambda f=filename: self.remove_file(f),
                bg='#f44336',
                fg='white',
                font=("Arial", 9, "bold"),
                padx=8,
                pady=8,
                corner_radius=6
            )
            remove_btn.pack(side='right', padx=(1, 0), pady=0)
            
            # Query by given categories button
            query_categories_btn = RoundedButton(
                file_frame,
                text="Query for Category(s)",
                command=lambda f=filename: self.show_category_dropdown(f),
                bg='#e3f2fd',
                fg='#1976d2',
                font=("Arial", 10),
                padx=15,
                pady=8,
                corner_radius=6
            )
            query_categories_btn.pack(side='right', padx=(1, 1), pady=0)
            
            # Query by keyword button
            query_keyword_btn = RoundedButton(
                file_frame,
                text="Query by Keyword",
                command=lambda f=filename: self.show_keyword_popup(f),
                bg='#e3f2fd',
                fg='#1976d2',
                font=("Arial", 10),
                padx=15,
                pady=8,
                corner_radius=6
            )
            query_keyword_btn.pack(side='right', padx=(1, 1), pady=0)
        
        # Update canvas scroll region
        self.canvas.update_idletasks()
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
    
    def update_output_list(self):
        """Update the output files display"""
        # Clear existing widgets in output scrollable frame
        for widget in self.output_scrollable_frame.winfo_children():
            widget.destroy()
        
        for i, (filename, content) in enumerate(self.output_files.items()):
            size_kb = len(content) / 1024
            if size_kb < 1024:
                size_str = f"{size_kb:.1f} KB"
            else:
                size_str = f"{size_kb/1024:.1f} MB"
            
            # Create frame for each output file entry
            output_frame = tk.Frame(self.output_scrollable_frame, bg='white', relief='flat', bd=0, highlightthickness=0)
            output_frame.pack(fill='x', padx=3, pady=1)
            
            # Output file info label (clickable to open file)
            output_info_label = RoundedButton(
                output_frame,
                text=f"📄 {filename} ({size_str})",
                command=lambda f=filename: self.open_output_file(f),
                bg='#f8f9fa',
                fg='#333333',
                hover_bg='#e9ecef',
                font=("Arial", 10),
                padx=12,
                pady=8,
                corner_radius=6
            )
            output_info_label.pack(side='left', fill='x', expand=True, padx=(0, 1), pady=0)
            
            # Red X button to remove output file
            remove_output_btn = RoundedButton(
                output_frame,
                text="✕",
                command=lambda f=filename: self.remove_output_file(f),
                bg='#f44336',
                fg='white',
                font=("Arial", 9, "bold"),
                padx=8,
                pady=8,
                corner_radius=6
            )
            remove_output_btn.pack(side='right', padx=(1, 0), pady=0)
        
        # Update canvas scroll region
        self.output_canvas.update_idletasks()
        self.output_canvas.configure(scrollregion=self.output_canvas.bbox("all"))

    def clear_all_files(self):
        """Clear all uploaded files"""
        if self.uploaded_files:
            result = messagebox.askyesno(
                "Confirm Clear", 
                "Are you sure you want to clear all uploaded files?"
            )
            if result:
                self.uploaded_files.clear()
                self.update_file_list()
                messagebox.showinfo("Cleared", "All uploaded files have been cleared.")
        else:
            messagebox.showinfo("No Files", "No uploaded files to clear.")
    
    def clear_all_outputs(self):
        """Clear all output files"""
        if self.output_files:
            result = messagebox.askyesno(
                "Confirm Clear", 
                "Are you sure you want to clear all output files?"
            )
            if result:
                self.output_files.clear()
                self.update_output_list()
                messagebox.showinfo("Cleared", "All output files have been cleared.")
        else:
            messagebox.showinfo("No Files", "No output files to clear.")

    def open_file(self, filename):
        """Open the uploaded file using the default system application"""
        import tempfile
        import subprocess
        import platform
        
        if filename not in self.uploaded_files:
            messagebox.showerror("Error", f"File '{filename}' not found.")
            return
        
        try:
            # Get file content
            file_content = self.uploaded_files[filename]['content']
            
            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            # Open the file with the default application
            if platform.system() == 'Windows':
                subprocess.run(['start', temp_file_path], shell=True, check=True)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.run(['open', temp_file_path], check=True)
            else:  # Linux
                subprocess.run(['xdg-open', temp_file_path], check=True)
                
            messagebox.showinfo("File Opened", f"Opening '{filename}' with default application.")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open file '{filename}':\n{str(e)}")

    def open_output_file(self, filename):
        """Open an output file using the default system application"""
        # Similar to open_file but for output files
        import tempfile
        import subprocess
        import platform
        
        if filename not in self.output_files:
            messagebox.showerror("Error", f"Output file '{filename}' not found.")
            return
        
        try:
            # Get file content
            file_content = self.output_files[filename]
            
            # Determine if content is binary (like DOCX) or text
            if isinstance(file_content, bytes):
                # Binary content (DOCX files)
                with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as temp_file:
                    temp_file.write(file_content)
                    temp_file_path = temp_file.name
            else:
                # Text content (TXT files)
                with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix, mode='w', encoding='utf-8') as temp_file:
                    temp_file.write(file_content)
                    temp_file_path = temp_file.name
            
            # Open the file with the default application
            if platform.system() == 'Windows':
                subprocess.run(['start', temp_file_path], shell=True, check=True)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.run(['open', temp_file_path], check=True)
            else:  # Linux
                subprocess.run(['xdg-open', temp_file_path], check=True)
                
            messagebox.showinfo("File Opened", f"Opening output file '{filename}' with default application.")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open output file '{filename}':\n{str(e)}")

    def remove_file(self, filename):
        """Remove a specific file from the uploaded files"""
        if filename in self.uploaded_files:
            result = messagebox.askyesno(
                "Confirm Remove", 
                f"Are you sure you want to remove '{filename}'?"
            )
            if result:
                del self.uploaded_files[filename]
                self.update_file_list()
                #messagebox.showinfo("Removed", f"File '{filename}' has been removed.")
        else:
            messagebox.showerror("Error", f"File '{filename}' not found.")

    def remove_output_file(self, filename):
        """Remove a specific output file"""
        if filename in self.output_files:
            result = messagebox.askyesno(
                "Confirm Remove", 
                f"Are you sure you want to remove output file '{filename}'?"
            )
            if result:
                del self.output_files[filename]
                self.update_output_list()
        else:
            messagebox.showerror("Error", f"Output file '{filename}' not found.")

    def open_policy_pdf(self):
        """Open the WFM policy PDF file"""
        import subprocess
        import platform
        
        try:
            # Get the directory where this script/executable is located
            if is_pyinstaller_bundle():
                # When running as PyInstaller EXE, get the directory of the executable
                script_dir = Path(sys.executable).parent
            else:
                # When running as script, get the script's directory
                script_dir = Path(__file__).parent
            
            # Construct path to WFMpolicy.pdf in the same folder
            policy_pdf_path = script_dir / "WFMpolicy.pdf"
            
            # Check if the file exists
            if not policy_pdf_path.exists():
                messagebox.showerror(
                    "File Not Found", 
                    f"WFMpolicy.pdf not found in the application folder.\n\nExpected location: {policy_pdf_path}"
                )
                return
            
            # Open the PDF with the default application
            if platform.system() == 'Windows':
                subprocess.run(['start', str(policy_pdf_path)], shell=True, check=True)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.run(['open', str(policy_pdf_path)], check=True)
            else:  # Linux
                subprocess.run(['xdg-open', str(policy_pdf_path)], check=True)
                
            messagebox.showinfo("Policy Opened", "Opening Work From Home Policy document.")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open policy document:\n{str(e)}")

    def upload_file(self):
        """Handle single file upload"""
        file_path = filedialog.askopenfilename(
            title="Select a file to upload",
            filetypes=[
                ("All Supported", "*.csv;*.pdf;*.doc;*.docx"),
                ("CSV files", "*.csv"),
                ("PDF files", "*.pdf"),
                ("Word files", "*.doc;*.docx"),
                ("All files", "*.*")
            ]
        )
        
        if file_path:
            self.process_file(file_path)
    
    def upload_folder(self):
        """Handle folder upload - processes all supported files in the selected folder"""
        import glob
        
        folder_path = filedialog.askdirectory(
            title="Select folder containing files to upload"
        )
        
        if folder_path:
            # Supported file extensions
            supported_extensions = ['*.csv', '*.pdf', '*.doc', '*.docx']
            
            # Find all supported files in the folder
            all_files = []
            for extension in supported_extensions:
                pattern = os.path.join(folder_path, extension)
                all_files.extend(glob.glob(pattern))
                # Also check for uppercase extensions
                pattern_upper = os.path.join(folder_path, extension.upper())
                all_files.extend(glob.glob(pattern_upper))
            
            if all_files:
                # Process each file found
                success_count = 0
                for file_path in all_files:
                    try:
                        self.process_file(file_path)
                        success_count += 1
                    except Exception as e:
                        print(f"Error processing {file_path}: {e}")
                

                
            else:
                messagebox.showinfo(
                    "No Files Found", 
                    f"No supported files (CSV, PDF, DOC, DOCX) found in the selected folder.\n\n"
                    f"Folder: {folder_path}"
                )
    
    def process_file(self, file_path):
        """Process and store a single file"""
        try:
            # Read file into variable
            file_content = read_file_as_variable(file_path)
            file_name = Path(file_path).name
            
            # Store both content and path for later use
            self.uploaded_files[file_name] = {
                'content': file_content,
                'path': file_path
            }
            
            # Update file list display
            self.update_file_list()
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to upload file:\n{str(e)}")
    
    def make_summary(self):
        """Create comprehensive summary by processing files in batches of 3, then creating master summary"""
        if not self.uploaded_files:
            messagebox.showerror("Error", "No files uploaded. Please upload files first.")
            return
        
        # Start the batch summary generation in a separate thread
        summary_thread = threading.Thread(target=self.execute_batch_summary_analysis)
        summary_thread.daemon = True
        summary_thread.start()
    
    def execute_batch_summary_analysis(self):
        """Execute batch summary analysis - 3 files at a time, then master summary"""
        try:
            # Prepare file list
            file_list = list(self.uploaded_files.keys())
            total_files = len(file_list)
            
            if total_files == 0:
                self.root.after(0, lambda: messagebox.showerror("Error", "No file content available for analysis."))
                return
            
            # Create batches with size and file count constraints
            batch_size_limit = 102 * 1024  # 102KB in bytes
            max_files_per_batch = 2
            
            batches = []
            current_batch = []
            current_batch_size = 0
            
            for filename in file_list:
                file_data = self.uploaded_files[filename]
                file_content = file_data.get('content', '')
                
                # Handle both string and bytes content for size calculation
                if isinstance(file_content, bytes):
                    file_size = len(file_content)
                else:
                    file_size = len(file_content.encode('utf-8'))  # Get size in bytes
                
                # Check if adding this file would exceed limits
                would_exceed_size = (current_batch_size + file_size) > batch_size_limit
                would_exceed_count = len(current_batch) >= max_files_per_batch
                
                # If adding this file would exceed limits, start a new batch
                if current_batch and (would_exceed_size or would_exceed_count):
                    batches.append(current_batch)
                    current_batch = []
                    current_batch_size = 0
                
                # Add file to current batch
                current_batch.append(filename)
                current_batch_size += file_size
                
                # If this single file exceeds size limit, create a batch with just this file
                if file_size > batch_size_limit:
                    batches.append(current_batch)
                    current_batch = []
                    current_batch_size = 0
            
            # Add the last batch if it has files
            if current_batch:
                batches.append(current_batch)
            
            batch_summaries = []
            
            # Process each batch
            for batch_num, batch_files in enumerate(batches, 1):
                try:
                    # Calculate batch size for display
                    batch_content_size = 0
                    for f in batch_files:
                        content = self.uploaded_files[f].get('content', '')
                        if isinstance(content, bytes):
                            batch_content_size += len(content)
                        else:
                            batch_content_size += len(content.encode('utf-8'))
                    batch_size_kb = batch_content_size / 1024
                    
                    # Create timer dialog for this batch
                    timer_dialog = TimerDialog(
                        self.root, 
                        f"Processing Batch {batch_num}/{len(batches)}", 
                        f"Analyzing {len(batch_files)} files ({batch_size_kb:.1f}KB): {', '.join(batch_files[:2])}{'...' if len(batch_files) > 2 else ''}"
                    )
                    dialog_window = timer_dialog.show()
                    self.root.update()
                    
                    # Prepare batch content
                    batch_content = []
                    for filename in batch_files:
                        file_data = self.uploaded_files[filename]
                        content = file_data.get('content', '')
                        if content:
                            batch_content.append(f"=== FILE: {filename} ===\n{content}\n")
                    
                    combined_batch_content = "\n".join(batch_content)
                    
                    # Prepare batch analysis prompt
                    batch_prompt = f"""
Please analyze the following {len(batch_files)} files for workforce management (WFM) policy compliance.

FILES IN THIS BATCH:
{', '.join(batch_files)}

ANALYSIS REQUIREMENTS:
1. POLICY COMPLIANCE FINDINGS
   - Key compliance issues identified
   - Areas of strong compliance
   - Critical violations or gaps
   
2. CATEGORIZED FINDINGS
   - Attendance and Time Management
   - Leave Policies and Procedures
   - Performance Management
   - Documentation Standards
   - Communication Protocols
   - Safety and Security Measures
   
3. RISK ASSESSMENT
   - High-risk issues requiring immediate attention
   - Medium-risk areas needing improvement
   - Low-risk observations
   
4. BATCH SUMMARY
   - Overall compliance status for these {len(batch_files)} files
   - Key recommendations for this batch
   - Priority actions needed

Please provide a focused analysis of these specific files:

{combined_batch_content}
"""
                    
                    # Process batch with AI
                    ai = PerplexityAI()
                    result = {'response': None, 'error': None, 'completed': False}
                    
                    def query_batch_ai():
                        try:
                            response = ai.query(batch_prompt)
                            result['response'] = response
                            result['completed'] = True
                        except Exception as e:
                            result['error'] = str(e)
                            result['completed'] = True
                    
                    # Start AI query in separate thread
                    ai_thread = threading.Thread(target=query_batch_ai)
                    ai_thread.daemon = True
                    ai_thread.start()
                    
                    # Wait for completion while keeping UI responsive
                    while not result['completed']:
                        self.root.update()
                        time.sleep(0.1)
                    
                    timer_dialog.close()
                    
                    if result['error']:
                        self.root.after(0, lambda e=result['error'], bn=batch_num: messagebox.showerror("API Error", 
                            f"Failed to analyze batch {bn}: {e}\n\nStopping batch processing."))
                        return
                    
                    # Save batch summary (for internal processing only, not displayed)
                    batch_filename = f"Batch_{batch_num}_Summary_{len(batch_files)}_files.txt"
                    # Don't add to output_files - keep only for master summary creation
                    batch_summaries.append({
                        'filename': batch_filename,
                        'content': result['response'],
                        'files': batch_files
                    })
                    
                    # Don't update display after each batch - only show final result
                    
                except Exception as batch_error:
                    if 'timer_dialog' in locals():
                        timer_dialog.close()
                    self.root.after(0, lambda e=str(batch_error), bn=batch_num: messagebox.showerror("Batch Error", 
                        f"Error processing batch {bn}: {e}"))
                    continue
            
            # Now create master summary from all batch summaries
            if batch_summaries:
                self.create_master_summary(batch_summaries, total_files)
            else:
                self.root.after(0, lambda: messagebox.showerror("Error", "No batch summaries were created successfully."))
                
        except Exception as e:
            error_msg = f"Batch summary analysis failed: {str(e)}"
            self.root.after(0, lambda: messagebox.showerror("Error", error_msg))
    
    def generate_summary_filename(self, extension="docx"):
        """Generate incremental summary filename (Summary.docx, Summary1.docx, Summary2.docx, etc.)"""
        base_name = "Summary"
        
        # Check if base name exists
        if f"{base_name}.{extension}" not in self.output_files:
            return f"{base_name}.{extension}"
        
        # Find the next available number
        counter = 1
        while f"{base_name}{counter}.{extension}" in self.output_files:
            counter += 1
        
        return f"{base_name}{counter}.{extension}"
    
    def create_master_summary(self, batch_summaries, total_files):
        """Create master summary from all batch summaries"""
        try:
            # Create timer dialog for master summary
            timer_dialog = TimerDialog(
                self.root, 
                "Creating Master Summary", 
                f"Combining {len(batch_summaries)} batch summaries into comprehensive report..."
            )
            dialog_window = timer_dialog.show()
            self.root.update()
            
            # Prepare master summary content
            batch_contents = []
            all_files_list = []
            
            for i, batch_info in enumerate(batch_summaries, 1):
                batch_contents.append(f"=== BATCH {i} SUMMARY ({len(batch_info['files'])} files: {', '.join(batch_info['files'])}) ===\n{batch_info['content']}\n")
                all_files_list.extend(batch_info['files'])
            
            combined_summaries = "\n".join(batch_contents)
            
            # Prepare master analysis prompt
            master_prompt = f"""
Please create a COMPREHENSIVE MASTER SUMMARY by analyzing the following {len(batch_summaries)} batch summaries covering {total_files} total files.

MASTER ANALYSIS REQUIREMENTS:
1. EXECUTIVE SUMMARY
   - Overall compliance status across ALL {total_files} files
   - Critical findings that span multiple batches
   - Organization-wide compliance score (1-10)

2. CONSOLIDATED FINDINGS BY CATEGORY
   - Attendance and Time Management (consolidated across all files)
   - Leave Policies and Procedures (organization-wide view)
   - Performance Management (systemic issues)
   - Documentation Standards (consistency analysis)
   - Communication Protocols (enterprise-wide assessment)
   - Safety and Security Measures (comprehensive review)

3. STRATEGIC RISK ASSESSMENT
   - Enterprise-level high-risk issues requiring immediate attention
   - Systemic medium-risk areas needing organizational improvement
   - Low-risk areas for optimization across the organization

4. COMPREHENSIVE RECOMMENDATIONS
   - Priority 1: Critical actions for immediate implementation
   - Priority 2: Medium-term organizational improvements
   - Priority 3: Long-term optimization strategies
   - Training and development needs
   - Policy revision requirements

5. IMPLEMENTATION ROADMAP
   - 30-day action items
   - 90-day improvement milestones
   - Annual policy review recommendations
   - Resource allocation suggestions

6. MASTER SCORECARD
   - Overall organizational compliance rating
   - Department/area-specific strengths and weaknesses
   - ROI projections for recommended improvements

FILES ANALYZED: {', '.join(all_files_list)}

BATCH SUMMARIES TO CONSOLIDATE:

{combined_summaries}
"""
            
            # Process master summary with AI
            ai = PerplexityAI()
            result = {'response': None, 'error': None, 'completed': False}
            
            def query_master_ai():
                try:
                    response = ai.query(master_prompt)
                    result['response'] = response
                    result['completed'] = True
                except Exception as e:
                    result['error'] = str(e)
                    result['completed'] = True
            
            # Start AI query in separate thread
            ai_thread = threading.Thread(target=query_master_ai)
            ai_thread.daemon = True
            ai_thread.start()
            
            # Wait for completion while keeping UI responsive
            while not result['completed']:
                self.root.update()
                time.sleep(0.1)
            
            timer_dialog.close()
            
            if result['error']:
                self.root.after(0, lambda: messagebox.showerror("API Error", 
                    f"Failed to create master summary: {result['error']}\n\n"
                    "Batch summaries are available, but master summary could not be created."))
                return
            
            # Create DOCX document for master summary
            docx_created = False
            master_filename = ""
            
            # Try to create DOCX document
            try:
                import docx
                from docx.shared import Inches
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                
                # Create a new document
                doc = docx.Document()
                
                # Add title
                title = doc.add_heading('COMPREHENSIVE WFM AUDIT SUMMARY', 0)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Add subtitle
                subtitle = doc.add_heading(f'Analysis of {total_files} Files in {len(batch_summaries)} Batches', level=2)
                subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Add master summary content only (no batch summaries)
                doc.add_heading('EXECUTIVE SUMMARY & STRATEGIC ANALYSIS', level=1)
                
                # Split the master summary into paragraphs and add them
                master_paragraphs = result['response'].split('\n\n')
                for paragraph_text in master_paragraphs:
                    if paragraph_text.strip():
                        paragraph = doc.add_paragraph(paragraph_text.strip())
                
                # Save the document to a bytes buffer
                import io
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_content = doc_buffer.getvalue()
                
                # Generate incremental summary filename
                master_filename = self.generate_summary_filename("docx")
                self.output_files[master_filename] = doc_content
                docx_created = True
                
            except ImportError:
                # Install python-docx and try again
                if self.install_package('python-docx'):
                    try:
                        # Retry after installation
                        import docx
                        from docx.shared import Inches
                        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                        
                        # Create a new document
                        doc = docx.Document()
                        
                        # Add title
                        title = doc.add_heading('COMPREHENSIVE WFM AUDIT SUMMARY', 0)
                        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        # Add subtitle
                        subtitle = doc.add_heading(f'Analysis of {total_files} Files in {len(batch_summaries)} Batches', level=2)
                        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        # Add master summary content only (no batch summaries)
                        doc.add_heading('EXECUTIVE SUMMARY & STRATEGIC ANALYSIS', level=1)
                        
                        # Split the master summary into paragraphs and add them
                        master_paragraphs = result['response'].split('\n\n')
                        for paragraph_text in master_paragraphs:
                            if paragraph_text.strip():
                                paragraph = doc.add_paragraph(paragraph_text.strip())
                        
                        # Save the document to a bytes buffer
                        import io
                        doc_buffer = io.BytesIO()
                        doc.save(doc_buffer)
                        doc_content = doc_buffer.getvalue()
                        
                        # Generate incremental summary filename
                        master_filename = self.generate_summary_filename("docx")
                        self.output_files[master_filename] = doc_content
                        docx_created = True
                        
                    except Exception as retry_error:
                        print(f"Failed to create DOCX after installation: {retry_error}")
                        docx_created = False
                
            # Fallback to text file if DOCX creation failed
            if not docx_created:
                master_content = result['response']
                master_filename = self.generate_summary_filename("txt")
                self.output_files[master_filename] = master_content
            
            # Update display
            self.root.after(0, self.update_output_list)
            self.root.after(0, lambda: messagebox.showinfo("Success", 
                f"Comprehensive batch analysis completed!\n\n"
                f"✅ Processed {total_files} files in {len(batch_summaries)} batches\n"
                f"✅ Each batch: ≤102KB, ≤2 files per batch\n"
                f"✅ Generated comprehensive master summary\n"
                f"✅ Formatted as professional Word document (.docx)\n"
                f"✅ Contains only the final consolidated analysis\n\n"
                f"Master summary saved as:\n{master_filename}\n\n"
                f"The Word document contains:\n"
                f"• Professional formatting with headings\n"
                f"• Executive summary and strategic analysis\n"
                f"• Consolidated findings from all {total_files} files\n"
                f"• Clean, focused final report without batch details"))
            
        except Exception as e:
            if 'timer_dialog' in locals():
                timer_dialog.close()
            error_msg = f"Master summary creation failed: {str(e)}"
            self.root.after(0, lambda: messagebox.showerror("Error", error_msg))
    
    def execute_summary_analysis(self):
        """Legacy function - kept for compatibility but redirects to batch processing"""
        self.execute_batch_summary_analysis()
    
    def show_category_dropdown(self, filename):
        """Show category dropdown checklist for the selected file"""
        # Create popup window
        category_window = tk.Toplevel(self.root)
        category_window.title(f"Query Categories - {filename}")
        category_window.configure(bg='#f0f0f0')
        category_window.resizable(False, False)
        
        # Center the window
        category_window.transient(self.root)
        category_window.grab_set()
        
        # Title label (smaller font)
        title_label = tk.Label(
            category_window,
            text=f"Select Categories for: {filename}",
            font=("Arial", 10, "bold"),
            bg='#f0f0f0',
            fg='#333333',
            wraplength=350
        )
        title_label.pack(pady=10, padx=10)
        
        # Frame for checkboxes
        checkbox_frame = tk.Frame(category_window, bg='#f0f0f0')
        checkbox_frame.pack(pady=10, padx=20, fill='both', expand=True)
        
        # Category options
        categories = [
            "Generic search based on all WFM policies",
            "Employee Eligibility",
            "Remote Work Agreements", 
            "IT Security Compliance",
            "Performance Monitoring",
            "Cybersecurity Training"
        ]
        
        # Dictionary to store checkbox variables
        self.category_vars = {}
        
        # Create checkboxes for each category
        for i, category in enumerate(categories):
            var = tk.BooleanVar()
            self.category_vars[category] = var
            
            # Special handling for the "Generic search" option
            if category == "Generic search based on all WFM policies":
                checkbox = tk.Checkbutton(
                    checkbox_frame,
                    text=category,
                    variable=var,
                    command=lambda: self.toggle_all_categories(),
                    font=("Arial", 10, "bold"),
                    bg='#f0f0f0',
                    fg='#000000',
                    selectcolor='#ffffff',
                    activebackground='#f0f0f0',
                    activeforeground='#000000',
                    relief='flat',
                    bd=0,
                    highlightthickness=0
                )
                checkbox.pack(anchor='w', pady=(5, 10), padx=10)
                
                # Add a separator line
                separator = tk.Frame(checkbox_frame, height=1, bg='#cccccc')
                separator.pack(fill='x', pady=(0, 10), padx=20)
            else:
                checkbox = tk.Checkbutton(
                    checkbox_frame,
                    text=category,
                    variable=var,
                    font=("Arial", 10),
                    bg='#f0f0f0',
                    fg='#333333',
                    selectcolor='#ffffff',
                    activebackground='#f0f0f0',
                    activeforeground='#333333',
                    relief='flat',
                    bd=0,
                    highlightthickness=0
                )
                checkbox.pack(anchor='w', pady=5, padx=10)
        
        # Buttons frame
        button_frame = tk.Frame(category_window, bg='#f0f0f0')
        button_frame.pack(pady=20)
        
        # Search button
        search_btn = RoundedButton(
            button_frame,
            text="Search",
            command=lambda: self.execute_category_search(filename, category_window),
            bg='#4CAF50',
            fg='white',
            font=("Arial", 10, "bold"),
            padx=20,
            pady=8,
            corner_radius=10
        )
        search_btn.pack(side='left', padx=10)
        
        # Cancel button
        cancel_btn = RoundedButton(
            button_frame,
            text="Cancel",
            command=category_window.destroy,
            bg='#757575',
            fg='white',
            font=("Arial", 10),
            padx=20,
            pady=8,
            corner_radius=10
        )
        cancel_btn.pack(side='left', padx=10)
        
        # Update window to get proper size, then adjust geometry
        category_window.update_idletasks()
        
        # Calculate required width based on content
        min_width = max(400, title_label.winfo_reqwidth() + 40)
        
        # Calculate required height dynamically based on actual content
        required_height = (
            title_label.winfo_reqheight() +  # Title
            checkbox_frame.winfo_reqheight() +  # All checkboxes
            button_frame.winfo_reqheight() +  # Buttons
            80  # Additional padding (top/bottom margins + extra space)
        )
        height = max(400, required_height)  # Minimum 400px, or calculated height
        
        # Set the final geometry
        category_window.geometry(f"{min_width}x{height}")
        
        # Center the window on screen
        x = (category_window.winfo_screenwidth() // 2) - (min_width // 2)
        y = (category_window.winfo_screenheight() // 2) - (height // 2)
        category_window.geometry(f"{min_width}x{height}+{x}+{y}")
    
    def toggle_all_categories(self):
        """Toggle all category checkboxes based on the 'Generic search' checkbox state"""
        # Get the state of the "Generic search" checkbox
        generic_search_var = self.category_vars.get("Generic search based on all WFM policies")
        if generic_search_var:
            is_checked = generic_search_var.get()
            
            # Set all other categories to the same state, excluding the generic search itself
            for category, var in self.category_vars.items():
                if category != "Generic search based on all WFM policies":
                    var.set(is_checked)
    
    def execute_category_search(self, filename, window):
        """Execute the category search using Perplexity AI with file analysis"""
        # Get selected categories (excluding the generic search toggle)
        selected_categories = [
            category for category, var in self.category_vars.items() 
            if var.get() and category != "Generic search based on all WFM policies"
        ]
        
        if not selected_categories:
            messagebox.showwarning("No Selection", "Please select at least one category.")
            return
        
        # Close the category window
        window.destroy()
        
        try:
            # Read file contents
            selected_file_content = self.read_file_content(filename)
            policy_file_content = self.read_policy_file()
            
            if not selected_file_content or not policy_file_content:
                return  # Error messages already shown in read functions
            
            # Create AI query
            categories_text = ", ".join(selected_categories)
            
            # Create the prompt for Perplexity AI
            prompt = f"""
            IMPORTANT: Use ONLY the two files provided below. Do NOT use web search or any external data sources.
            
            I need you to analyze these two files for information relevant to the following categories: {categories_text}
            
            FILE 1 - Selected File ({filename}):
            {selected_file_content}
            
            FILE 2 - WFM Policy Document (WFMpolicy.pdf):
            {policy_file_content}
            
            Please analyze both files and:
            1. Identify where the WFM policy requirements are being met for the selected categories
            2. Identify where the WFM policy requirements are NOT being met for the selected categories
            3. If the selected file does not contain information about the selected categories, state: "This file does not have any information regarding {categories_text}" and provide a brief summary of what other policies/topics it might address instead.
            
            Provide a clear, structured summary of your findings. Use ONLY the content from these two files - no external sources.
            """
            
            # Send query to Perplexity AI with search disabled
            ai = PerplexityAI()
            
            # Show timer dialog
            timer_dialog = TimerDialog(self.root, "Processing Category Search", "Analyzing files with AI for categories...")
            dialog_window = timer_dialog.show()
            
            # Force the dialog to appear
            self.root.update()
            
            # Variables to store result and error
            result = {'response': None, 'error': None, 'completed': False}
            
            def ai_query_thread():
                try:
                    result['response'] = ai.query(prompt, disable_search=True)
                except Exception as e:
                    result['error'] = str(e)
                finally:
                    result['completed'] = True
            
            # Start AI query in background thread
            thread = threading.Thread(target=ai_query_thread)
            thread.daemon = True
            thread.start()
            
            # Wait for completion while updating the dialog
            while not result['completed']:
                self.root.update()
                self.root.after(100)  # Wait 100ms before checking again
            
            # Close timer dialog
            timer_dialog.close()
            
            # Check for errors
            if result['error']:
                raise Exception(result['error'])
            
            response = result['response']
            
            # Create DOCX output instead of popup
            self.create_search_result_docx(filename, "Category Search", categories_text, response)
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to analyze files:\n{str(e)}")
    
    def create_search_result_docx(self, filename, search_type, search_criteria, response):
        """Create a DOCX file for search results and add to output files"""
        try:
            # Try to create DOCX document
            docx_created = False
            
            try:
                import docx
                from docx.shared import Inches
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                
                # Create a new document
                doc = docx.Document()
                
                # Add title
                title = doc.add_heading(f'{search_type} Results', 0)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Add subtitle with file name
                subtitle = doc.add_heading(f'Analysis of: {filename}', level=2)
                subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Add search criteria
                doc.add_heading('Search Criteria', level=1)
                doc.add_paragraph(search_criteria)
                
                # Add analysis results
                doc.add_heading('Analysis Results', level=1)
                
                # Split the response into paragraphs and add them
                response_paragraphs = response.split('\n\n')
                for paragraph_text in response_paragraphs:
                    if paragraph_text.strip():
                        doc.add_paragraph(paragraph_text.strip())
                
                # Save the document to a bytes buffer
                import io
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_content = doc_buffer.getvalue()
                
                # Generate filename with timestamp
                import datetime
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                result_filename = f"{search_type.replace(' ', '_')}_{filename}_{timestamp}.docx"
                
                # Add to output files
                self.output_files[result_filename] = doc_content
                docx_created = True
                
            except ImportError:
                # Install python-docx and try again
                if self.install_package('python-docx'):
                    try:
                        # Retry after installation
                        import docx
                        from docx.shared import Inches
                        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                        
                        # Create a new document
                        doc = docx.Document()
                        
                        # Add title
                        title = doc.add_heading(f'{search_type} Results', 0)
                        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        # Add subtitle with file name
                        subtitle = doc.add_heading(f'Analysis of: {filename}', level=2)
                        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        # Add search criteria
                        doc.add_heading('Search Criteria', level=1)
                        doc.add_paragraph(search_criteria)
                        
                        # Add analysis results
                        doc.add_heading('Analysis Results', level=1)
                        
                        # Split the response into paragraphs and add them
                        response_paragraphs = response.split('\n\n')
                        for paragraph_text in response_paragraphs:
                            if paragraph_text.strip():
                                doc.add_paragraph(paragraph_text.strip())
                        
                        # Save the document to a bytes buffer
                        import io
                        doc_buffer = io.BytesIO()
                        doc.save(doc_buffer)
                        doc_content = doc_buffer.getvalue()
                        
                        # Generate filename with timestamp
                        import datetime
                        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                        result_filename = f"{search_type.replace(' ', '_')}_{filename}_{timestamp}.docx"
                        
                        # Add to output files
                        self.output_files[result_filename] = doc_content
                        docx_created = True
                        
                    except Exception as retry_error:
                        print(f"Failed to create DOCX after installation: {retry_error}")
                        docx_created = False
            
            # Fallback to text file if DOCX creation failed
            if not docx_created:
                import datetime
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                result_filename = f"{search_type.replace(' ', '_')}_{filename}_{timestamp}.txt"
                
                # Create text content
                text_content = f"{search_type} Results\n"
                text_content += f"Analysis of: {filename}\n\n"
                text_content += f"Search Criteria: {search_criteria}\n\n"
                text_content += "Analysis Results:\n\n"
                text_content += response
                
                # Add to output files
                self.output_files[result_filename] = text_content
            
            # Update the output display and show success message
            self.update_output_list()
            messagebox.showinfo("Search Complete", 
                f"{search_type} completed successfully!\n\n"
                f"File analyzed: {filename}\n"
                f"Search criteria: {search_criteria}\n\n"
                f"Results saved as: {result_filename}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to create search result file: {str(e)}")
    
    def read_file_content(self, filename):
        """Read the content of the selected file, converting PDF/DOCX to text"""
        try:
            # Check if file exists in uploaded files
            if filename not in self.uploaded_files:
                messagebox.showerror("File Not Found", f"Could not find file: {filename}")
                return None
            
            # Get file data (both content and path)
            file_data = self.uploaded_files[filename]
            file_path = file_data['path']
            file_extension = Path(file_path).suffix.lower()
            
            # Handle different file types
            if file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            elif file_extension == '.csv':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            elif file_extension == '.pdf':
                return self.convert_pdf_to_text(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self.convert_docx_to_text(file_path)
            else:
                messagebox.showerror(
                    "Unsupported Format", 
                    f"File format {file_extension} is not supported for AI analysis.\n"
                    f"Supported formats: .txt, .csv, .pdf, .docx, .doc"
                )
                return None
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to read file {filename}:\n{str(e)}")
            return None
    
    def install_package(self, package_name):
        """Handle package installation - different behavior for PyInstaller vs development"""
        # Check if running as PyInstaller bundle
        if is_pyinstaller_bundle():
            # In PyInstaller bundle, packages should already be included
            messagebox.showerror(
                "Missing Dependency", 
                f"The required package '{package_name}' is missing from this application.\n\n"
                f"This appears to be a packaging issue. The application was built without "
                f"including the '{package_name}' dependency.\n\n"
                f"Please contact the application developer to rebuild with all dependencies included."
            )
            return False
        else:
            # Development environment - attempt pip install
            try:
                import subprocess
                import sys
                
                print(f"Installing {package_name}...")
                messagebox.showinfo("Installing Package", f"Installing {package_name}... Please wait.")
                
                # Run pip install
                result = subprocess.run(
                    [sys.executable, "-m", "pip", "install", package_name],
                    capture_output=True,
                    text=True,
                    check=True
                )
                
                print(f"Successfully installed {package_name}")
                messagebox.showinfo("Installation Complete", f"Successfully installed {package_name}")
                return True
                
            except subprocess.CalledProcessError as e:
                error_msg = f"Failed to install {package_name}:\n{e.stderr}"
                print(error_msg)
                messagebox.showerror("Installation Failed", error_msg)
                return False
            except Exception as e:
                error_msg = f"Error installing {package_name}: {str(e)}"
                print(error_msg)
                messagebox.showerror("Installation Error", error_msg)
                return False
    
    def convert_pdf_to_text(self, pdf_path):
        """Convert PDF file to text using built-in libraries"""
        try:
            # Try using PyPDF2 if available
            try:
                import PyPDF2
                text = ""
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                return text.strip()
            except ImportError:
                # Try to install PyPDF2 automatically
                if messagebox.askyesno("Install PDF Library", 
                                     "PyPDF2 library is required to read PDF files.\n\n"
                                     "Would you like to install it automatically?"):
                    if self.install_package("PyPDF2"):
                        # Try importing again after installation
                        try:
                            import PyPDF2
                            text = ""
                            with open(pdf_path, 'rb') as file:
                                pdf_reader = PyPDF2.PdfReader(file)
                                for page in pdf_reader.pages:
                                    text += page.extract_text() + "\n"
                            return text.strip()
                        except ImportError:
                            pass
            
            # Try using pdfplumber if PyPDF2 failed
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                return text.strip()
            except ImportError:
                # Try to install pdfplumber automatically
                if messagebox.askyesno("Install Alternative PDF Library", 
                                     "pdfplumber library could also work for PDF files.\n\n"
                                     "Would you like to install it instead?"):
                    if self.install_package("pdfplumber"):
                        # Try importing again after installation
                        try:
                            import pdfplumber
                            text = ""
                            with pdfplumber.open(pdf_path) as pdf:
                                for page in pdf.pages:
                                    page_text = page.extract_text()
                                    if page_text:
                                        text += page_text + "\n"
                            return text.strip()
                        except ImportError:
                            pass
            
            # If both automatic installations failed or were declined
            messagebox.showerror(
                "PDF Library Required",
                "Unable to read PDF files. Please install manually:\n\n"
                "pip install PyPDF2\n"
                "OR\n"
                "pip install pdfplumber\n\n"
                "Or convert the PDF to a .txt file manually."
            )
            return None
            
        except Exception as e:
            messagebox.showerror("PDF Error", f"Failed to convert PDF to text:\n{str(e)}")
            return None
    
    def convert_docx_to_text(self, docx_path):
        """Convert DOCX file to text"""
        try:
            # Try using python-docx if available
            try:
                import docx
                doc = docx.Document(docx_path)
                text = []
                for paragraph in doc.paragraphs:
                    text.append(paragraph.text)
                return "\n".join(text)
            except ImportError:
                # Try to install python-docx automatically
                if messagebox.askyesno("Install DOCX Library", 
                                     "python-docx library is required to read DOCX files.\n\n"
                                     "Would you like to install it automatically?"):
                    if self.install_package("python-docx"):
                        # Try importing again after installation
                        try:
                            import docx
                            doc = docx.Document(docx_path)
                            text = []
                            for paragraph in doc.paragraphs:
                                text.append(paragraph.text)
                            return "\n".join(text)
                        except ImportError:
                            pass
            
            # Try using zipfile to extract text (basic method for .docx)
            try:
                import zipfile
                import xml.etree.ElementTree as ET
                
                text = ""
                with zipfile.ZipFile(docx_path, 'r') as zip_file:
                    # Read the main document
                    with zip_file.open('word/document.xml') as xml_file:
                        tree = ET.parse(xml_file)
                        root = tree.getroot()
                        
                        # Extract text from all text nodes
                        namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                        for text_element in root.findall('.//w:t', namespace):
                            if text_element.text:
                                text += text_element.text
                        
                        # Add paragraph breaks
                        for para in root.findall('.//w:p', namespace):
                            text += "\n"
                
                return text.strip()
            except Exception:
                pass
            
            # If automatic installation failed or was declined
            messagebox.showerror(
                "DOCX Library Required",
                "Unable to read DOCX files. Please install manually:\n\n"
                "pip install python-docx\n\n"
                "Or convert the DOCX to a .txt file manually."
            )
            return None
            
        except Exception as e:
            messagebox.showerror("DOCX Error", f"Failed to convert DOCX to text:\n{str(e)}")
            return None
    
    def read_policy_file(self):
        """Read the WFM policy file content, converting PDF to text if needed"""
        try:
            # Get the directory where this script/executable is located
            if is_pyinstaller_bundle():
                # When running as PyInstaller EXE, get the directory of the executable
                script_dir = Path(sys.executable).parent
            else:
                # When running as script, get the script's directory
                script_dir = Path(__file__).parent
            
            # Try different file formats in order of preference
            policy_files = [
                ("WFMpolicy.txt", "txt"),
                ("WFMpolicy.pdf", "pdf"),
                ("WFMpolicy.docx", "docx"),
                ("WFMpolicy.doc", "doc")
            ]
            
            for filename, file_type in policy_files:
                policy_path = script_dir / filename
                if policy_path.exists():
                    if file_type == "txt":
                        with open(policy_path, 'r', encoding='utf-8') as f:
                            return f.read()
                    elif file_type == "pdf":
                        return self.convert_pdf_to_text(policy_path)
                    elif file_type in ["docx", "doc"]:
                        return self.convert_docx_to_text(policy_path)
            
            # If no policy file found
            messagebox.showerror(
                "Policy File Not Found", 
                f"WFM policy file not found in {script_dir}\n\n"
                f"Please ensure one of these files exists:\n"
                f"• WFMpolicy.txt (preferred)\n"
                f"• WFMpolicy.pdf\n"
                f"• WFMpolicy.docx\n"
                f"• WFMpolicy.doc"
            )
            return None
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to read policy file:\n{str(e)}")
            return None
    
    def show_ai_response(self, filename, categories, response):
        """Show AI response in a popup window"""
        # Create popup window
        response_window = tk.Toplevel(self.root)
        response_window.title(f"AI Analysis Results - {filename}")
        response_window.configure(bg='#f0f0f0')
        response_window.geometry("800x600")
        
        # Center the window
        response_window.transient(self.root)
        response_window.grab_set()
        
        # Title label
        title_label = tk.Label(
            response_window,
            text=f"AI Analysis for: {filename}",
            font=("Arial", 12, "bold"),
            bg='#f0f0f0',
            fg='#333333'
        )
        title_label.pack(pady=10, padx=10)
        
        # Categories label
        categories_label = tk.Label(
            response_window,
            text=f"Categories Analyzed: {categories}",
            font=("Arial", 10),
            bg='#f0f0f0',
            fg='#666666',
            wraplength=750
        )
        categories_label.pack(pady=(0, 10), padx=10)
        
        # Text frame with scrollbar
        text_frame = tk.Frame(response_window, bg='#f0f0f0')
        text_frame.pack(fill='both', expand=True, padx=20, pady=10)
        
        # Scrollable text widget
        text_widget = tk.Text(
            text_frame,
            wrap='word',
            font=("Arial", 10),
            bg='white',
            fg='#333333',
            padx=10,
            pady=10
        )
        
        # Scrollbar
        scrollbar = tk.Scrollbar(text_frame, orient="vertical", command=text_widget.yview)
        text_widget.configure(yscrollcommand=scrollbar.set)
        
        # Pack text widget and scrollbar
        text_widget.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
        # Insert AI response
        text_widget.insert('1.0', response)
        text_widget.config(state='disabled')  # Make read-only
        
        # Close button
        close_btn = RoundedButton(
            response_window,
            text="Close",
            command=response_window.destroy,
            bg='#757575',
            fg='white',
            font=("Arial", 10),
            padx=30,
            pady=8,
            corner_radius=10
        )
        close_btn.pack(pady=20)
    
    def show_keyword_popup(self, filename):
        """Show keyword search popup for the selected file"""
        # Create popup window
        keyword_window = tk.Toplevel(self.root)
        keyword_window.title(f"Query by Keyword - {filename}")
        keyword_window.configure(bg='#f0f0f0')
        keyword_window.resizable(False, False)
        
        # Center the window
        keyword_window.transient(self.root)
        keyword_window.grab_set()
        
        # Title label (smaller font)
        title_label = tk.Label(
            keyword_window,
            text=f"Enter Keywords for: {filename}",
            font=("Arial", 10, "bold"),
            bg='#f0f0f0',
            fg='#333333',
            wraplength=350
        )
        title_label.pack(pady=10, padx=10)
        
        # Instruction label
        instruction_label = tk.Label(
            keyword_window,
            text="Enter keywords separated by commas:",
            font=("Arial", 9),
            bg='#f0f0f0',
            fg='#666666'
        )
        instruction_label.pack(pady=(0, 10), padx=10)
        
        # Frame for textbox
        textbox_frame = tk.Frame(keyword_window, bg='#f0f0f0')
        textbox_frame.pack(pady=10, padx=20, fill='x')
        
        # Flat textbox for keywords
        self.keyword_entry = tk.Text(
            textbox_frame,
            height=4,
            width=40,
            font=("Arial", 10),
            bg='white',
            fg='#333333',
            relief='flat',
            bd=1,
            highlightthickness=1,
            highlightcolor='#2196F3',
            highlightbackground='#cccccc',
            wrap='word'
        )
        self.keyword_entry.pack(fill='x', pady=5)
        
        # Buttons frame
        button_frame = tk.Frame(keyword_window, bg='#f0f0f0')
        button_frame.pack(pady=20)
        
        # Search button
        search_btn = RoundedButton(
            button_frame,
            text="Search",
            command=lambda: self.execute_keyword_search(filename, keyword_window),
            bg='#4CAF50',
            fg='white',
            font=("Arial", 10, "bold"),
            padx=20,
            pady=8,
            corner_radius=10
        )
        search_btn.pack(side='left', padx=10)
        
        # Cancel button
        cancel_btn = RoundedButton(
            button_frame,
            text="Cancel",
            command=keyword_window.destroy,
            bg='#757575',
            fg='white',
            font=("Arial", 10),
            padx=20,
            pady=8,
            corner_radius=10
        )
        cancel_btn.pack(side='left', padx=10)
        
        # Update window to get proper size, then adjust geometry
        keyword_window.update_idletasks()
        
        # Calculate required width based on content
        min_width = max(400, title_label.winfo_reqwidth() + 40)
        
        # Calculate required height dynamically based on actual content
        required_height = (
            title_label.winfo_reqheight() +  # Title
            instruction_label.winfo_reqheight() +  # Instruction
            textbox_frame.winfo_reqheight() +  # Text box
            button_frame.winfo_reqheight() +  # Buttons
            80  # Additional padding (top/bottom margins + extra space)
        )
        height = max(320, required_height)  # Minimum 320px, or calculated height
        
        # Set the final geometry
        keyword_window.geometry(f"{min_width}x{height}")
        
        # Center the window on screen
        x = (keyword_window.winfo_screenwidth() // 2) - (min_width // 2)
        y = (keyword_window.winfo_screenheight() // 2) - (height // 2)
        keyword_window.geometry(f"{min_width}x{height}+{x}+{y}")
        
        # Focus on the text entry
        self.keyword_entry.focus()
    
    def execute_keyword_search(self, filename, window):
        """Execute the keyword search using Perplexity AI with file analysis"""
        # Get entered keywords
        keywords_text = self.keyword_entry.get("1.0", tk.END).strip()
        
        if not keywords_text:
            messagebox.showwarning("No Keywords", "Please enter at least one keyword.")
            return
        
        # Close the keyword window
        window.destroy()
        
        try:
            # Read file contents
            selected_file_content = self.read_file_content(filename)
            policy_file_content = self.read_policy_file()
            
            if not selected_file_content or not policy_file_content:
                return  # Error messages already shown in read functions
            
            # Clean up keywords for display
            keywords_list = [kw.strip() for kw in keywords_text.split(',') if kw.strip()]
            keywords_display = ", ".join(keywords_list)
            
            # Create the prompt for Perplexity AI
            prompt = f"""
            IMPORTANT: Use ONLY the two files provided below. Do NOT use web search or any external data sources.
            
            I need you to analyze these two files for information relevant to the following keywords: {keywords_display}
            
            FILE 1 - Selected File ({filename}):
            {selected_file_content}
            
            FILE 2 - WFM Policy Document (WFMpolicy.pdf):
            {policy_file_content}
            
            Please analyze both files and:
            1. Search for any mention or reference to the keywords: {keywords_display}
            2. Identify where the WFM policy requirements related to these keywords are being met
            3. Identify where the WFM policy requirements related to these keywords are NOT being met
            4. If the selected file does not contain information about the keywords, respond with: "This document does not address the keyword(s) '{keywords_display}'"
            
            Keep the response concise and focused on the keywords. Use ONLY the content from these two files - no external sources.
            """
            
            # Send query to Perplexity AI with search disabled
            ai = PerplexityAI()
            
            # Show timer dialog
            timer_dialog = TimerDialog(self.root, "Processing Keyword Search", "Analyzing files with AI for keywords...")
            dialog_window = timer_dialog.show()
            
            # Force the dialog to appear
            self.root.update()
            
            # Variables to store result and error
            result = {'response': None, 'error': None, 'completed': False}
            
            def ai_query_thread():
                try:
                    result['response'] = ai.query(prompt, disable_search=True)
                except Exception as e:
                    result['error'] = str(e)
                finally:
                    result['completed'] = True
            
            # Start AI query in background thread
            thread = threading.Thread(target=ai_query_thread)
            thread.daemon = True
            thread.start()
            
            # Wait for completion while updating the dialog
            while not result['completed']:
                self.root.update()
                self.root.after(100)  # Wait 100ms before checking again
            
            # Close timer dialog
            timer_dialog.close()
            
            # Check for errors
            if result['error']:
                raise Exception(result['error'])
            
            response = result['response']
            
            # Create DOCX output instead of popup
            self.create_search_result_docx(filename, "Keyword Search", f"Keywords: {keywords_display}", response)
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to analyze files:\n{str(e)}")
    
    def open_policy_pdf(self):
        """Open the WFM policy PDF file"""
        import subprocess
        import platform
        
        try:
            # Get the directory where this script/executable is located
            if is_pyinstaller_bundle():
                # When running as PyInstaller EXE, get the directory of the executable
                script_dir = Path(sys.executable).parent
            else:
                # When running as script, get the script's directory
                script_dir = Path(__file__).parent
            
            # Construct path to WFMpolicy.pdf in the same folder
            policy_pdf_path = script_dir / "WFMpolicy.pdf"
            
            # Check if the file exists
            if not policy_pdf_path.exists():
                messagebox.showerror(
                    "File Not Found", 
                    f"WFMpolicy.pdf not found in the application folder.\n\nExpected location: {policy_pdf_path}"
                )
                return
            
            # Open the PDF with the default application
            if platform.system() == 'Windows':
                subprocess.run(['start', str(policy_pdf_path)], shell=True, check=True)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.run(['open', str(policy_pdf_path)], check=True)
            else:  # Linux
                subprocess.run(['xdg-open', str(policy_pdf_path)], check=True)
                
            messagebox.showinfo("Policy Opened", "Opening Work From Home Policy document.")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open policy document:\n{str(e)}")

def run_gui():
    """Run the GUI application"""
    root = tk.Tk()
    app = FileUploaderGUI(root)
    root.mainloop()

# Run the GUI when script is executed
if __name__ == "__main__":
    run_gui()
